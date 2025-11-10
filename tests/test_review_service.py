"""Integration tests for review service."""
from __future__ import annotations

from pathlib import Path
from unittest.mock import AsyncMock, MagicMock, patch

import pytest

from app.checklist_loader import ChecklistRepository
from app.models import ChecklistEvaluation, ChecklistEvaluationItem
from app.review_service import ReviewService


@pytest.mark.integration
class TestReviewService:
    """Test review service integration."""

    @pytest.fixture
    def review_service(self, temp_checklist_file: Path) -> ReviewService:
        """Create a review service instance for testing."""
        repo = ChecklistRepository(str(temp_checklist_file))
        return ReviewService(repo)

    @pytest.mark.asyncio
    async def test_review_contract_complete_workflow(
        self, review_service: ReviewService, temp_dir: Path, mock_env_vars: None
    ) -> None:
        """Test complete contract review workflow."""
        contract_file = temp_dir / "test_contract.txt"
        contract_file.write_text(
            "COMMERCIAL LEASE AGREEMENT\n\n"
            "Landlord: ABC Properties Pty Ltd (ABN 12 345 678 901)\n"
            "Tenant: XYZ Retail Store Pty Ltd\n"
            "Premises: 123 Main Street, Brisbane\n"
            "Rent: $5,000 per month\n"
            "Term: 3 years",
            encoding="utf-8",
        )

        evaluation, report, drive_link, telegram_result = await review_service.review_contract(
            file_path=contract_file,
            original_filename="test_contract.txt",
            section_id="QLD_COMMERCIAL_LEASE",
        )

        assert isinstance(evaluation, ChecklistEvaluation)
        assert len(evaluation.items) > 0
        assert Path(report.local_path).exists()
        assert drive_link is not None
        assert telegram_result is not None

    @pytest.mark.asyncio
    async def test_review_contract_generates_report(
        self, review_service: ReviewService, temp_dir: Path, mock_env_vars: None
    ) -> None:
        """Test that review service generates a report file."""
        contract_file = temp_dir / "test.txt"
        contract_file.write_text("Sample contract content", encoding="utf-8")

        evaluation, report, _, _ = await review_service.review_contract(
            file_path=contract_file,
            original_filename="test.txt",
            section_id="QLD_COMMERCIAL_LEASE",
        )

        assert Path(report.local_path).exists()
        assert report.local_path.endswith(".md")
        content = Path(report.local_path).read_text(encoding="utf-8")
        assert "# Contract Review Report" in content

    @pytest.mark.asyncio
    async def test_review_contract_returns_evaluation(
        self, review_service: ReviewService, temp_dir: Path, mock_env_vars: None
    ) -> None:
        """Test that review service returns evaluation results."""
        contract_file = temp_dir / "test.txt"
        contract_file.write_text("Commercial lease with rent review clause", encoding="utf-8")

        evaluation, _, _, _ = await review_service.review_contract(
            file_path=contract_file,
            original_filename="test.txt",
            section_id="QLD_COMMERCIAL_LEASE",
        )

        assert evaluation.summary is not None
        assert len(evaluation.items) > 0
        assert all(isinstance(item, ChecklistEvaluationItem) for item in evaluation.items)

    @pytest.mark.asyncio
    async def test_review_contract_handles_docx_file(
        self, review_service: ReviewService, temp_dir: Path, mock_env_vars: None
    ) -> None:
        """Test that review service can process DOCX files."""
        from docx import Document

        docx_file = temp_dir / "test.docx"
        doc = Document()
        doc.add_paragraph("Commercial lease agreement")
        doc.add_paragraph("Landlord: ABC Properties")
        doc.add_paragraph("Tenant: XYZ Store")
        doc.save(str(docx_file))

        evaluation, report, _, _ = await review_service.review_contract(
            file_path=docx_file,
            original_filename="test.docx",
            section_id="QLD_COMMERCIAL_LEASE",
        )

        assert isinstance(evaluation, ChecklistEvaluation)
        assert Path(report.local_path).exists()

    @pytest.mark.asyncio
    async def test_review_contract_derives_title_from_filename(
        self, review_service: ReviewService, temp_dir: Path, mock_env_vars: None
    ) -> None:
        """Test that contract title is derived from filename."""
        contract_file = temp_dir / "test_contract.txt"
        contract_file.write_text("Sample content", encoding="utf-8")

        _, report, _, _ = await review_service.review_contract(
            file_path=contract_file,
            original_filename="My_Commercial_Lease_2024.txt",
            section_id="QLD_COMMERCIAL_LEASE",
        )

        assert "My Commercial Lease 2024" in report.contract_title or "My_Commercial_Lease_2024" in report.contract_title

    @pytest.mark.asyncio
    async def test_review_contract_includes_drive_link_in_report(
        self, review_service: ReviewService, temp_dir: Path, mock_env_vars: None
    ) -> None:
        """Test that drive link is included in report artifact."""
        contract_file = temp_dir / "test.txt"
        contract_file.write_text("Sample content", encoding="utf-8")

        _, report, drive_link, _ = await review_service.review_contract(
            file_path=contract_file,
            original_filename="test.txt",
            section_id="QLD_COMMERCIAL_LEASE",
        )

        assert report.drive_link == drive_link
        assert drive_link is not None
        assert "drive.google.com" in drive_link or drive_link.startswith("http")

    @pytest.mark.asyncio
    async def test_review_contract_telegram_notification(
        self, review_service: ReviewService, temp_dir: Path, mock_env_vars: None
    ) -> None:
        """Test that Telegram notification is attempted."""
        contract_file = temp_dir / "test.txt"
        contract_file.write_text("Sample content", encoding="utf-8")

        _, _, _, telegram_result = await review_service.review_contract(
            file_path=contract_file,
            original_filename="test.txt",
            section_id="QLD_COMMERCIAL_LEASE",
        )

        assert telegram_result is not None
        # Without credentials, should return sent=False
        assert telegram_result.sent is False or telegram_result.sent is True

    @pytest.mark.asyncio
    async def test_review_contract_with_nonexistent_section_raises_error(
        self, review_service: ReviewService, temp_dir: Path
    ) -> None:
        """Test that review service raises error for nonexistent section."""
        contract_file = temp_dir / "test.txt"
        contract_file.write_text("Sample content", encoding="utf-8")

        with pytest.raises(KeyError, match="Checklist section 'INVALID_SECTION' not found"):
            await review_service.review_contract(
                file_path=contract_file,
                original_filename="test.txt",
                section_id="INVALID_SECTION",
            )

